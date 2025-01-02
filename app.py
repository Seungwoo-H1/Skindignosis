# app.py
import streamlit as st
from PIL import Image
from utils.CustomVisionModel import CustomVisionModel
from utils.PhiModel import PhiModel
from utils.LlamaModel import LlamaModel
from utils.CustomVisionHandler import CustomVisionHandler
from utils.FeedbackHandler import FeedbackHandler
from utils.TranslationHandler import TranslationHandler  # 번역 클래스 임포트
import deepl
from io import BytesIO
import base64
from docx import Document

# Setting Streamlit page
# Streamlit 페이지 설정
st.set_page_config(page_title="Know Your Skin", layout="centered")

# Background image path
# 배경 이미지 파일 경로 설정
background_image_path = "data/background.jpg"

# Function to encode image to Base64
# 이미지를 Base64로 인코딩하는 함수 정의
def get_base64_image(image_path):
    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode()
    return encoded_string

# Get Base64 encoded image data
# Base64로 인코딩된 이미지 데이터 가져오기
base64_background_image = get_base64_image(background_image_path)

# HTML and CSS for background image and animation
# 배경 이미지 및 애니메이션 텍스트 설정을 위한 HTML 및 CSS
st.markdown(
    f"""
    <style>
    /* Setting background image for the whole page */
    /* 전체 페이지 배경 이미지 설정 */
    .stApp {{
        background-image: url("data:image/jpeg;base64,{base64_background_image}");
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
        background-position: center;
    }}
    
    /* Animation text container */
    /* 애니메이션 텍스트 컨테이너 */
    .container {{
        overflow: hidden;
        white-space: nowrap;
        width: 100%;
        text-align: center;
        margin-top: 50px;
        padding: 20px;
        background-color: rgba(255, 255, 255, 0.6); /* Semi-transparent background for readability */
        /* 가독성을 위한 반투명 배경 */
        border-radius: 10px;
    }}

    /* Animation text */
    /* 애니메이션 텍스트 */
    .sliding-text {{
        display: inline-block;
        font-size: 3em;
        font-weight: bold;
        color: #ff6347; /* Text color */
        /* 텍스트 색상 */
        animation: slide 7s linear infinite;
    }}

    /* Slide animation */
    /* 슬라이드 애니메이션 */
    @keyframes slide {{
        0% {{
            transform: translateX(-100%);
        }}
        100% {{
            transform: translateX(100%);
        }}
    }}
    </style>

    <!-- HTML structure containing animated text -->
    <!-- 애니메이션 텍스트를 포함하는 HTML 구조 -->
    <div class="container">
        <div class="sliding-text">Know Your Skin</div>
    </div>
    """,
    unsafe_allow_html=True
)

# Setting Azure Custom Vision and models
# Azure Custom Vision 및 모델 설정
PREDICTION_KEY = "Enter your prediction key here"
ENDPOINT = "Enter your endpoint here"
PROJECT_ID = "Enter your project ID here"
ITERATION_NAME = "Iteration1"

MODEL_CATALOG_ENDPOINT = "Enter your model catalog endpoint here"
MODEL_CATALOG_API_KEY = "Enter your model catalog API key here"

LLAMA_TOKEN = "Enter your llama token here"
LLAMA_ENDPOINT = "Enter your llama endpoint here"
LLAMA_MODEL_NAME = "Meta-Llama-3.1-405B-Instruct"

DEEPL_API_KEY = "Enter your deepl API key here"

# Initializing translation handler instance
# 번역 핸들러 인스턴스 초기화
translation_handler = TranslationHandler(DEEPL_API_KEY)

# Initializing Custom Vision, Phi, and Llama model instances
# Custom Vision, Phi 및 Llama 모델 인스턴스 초기화
custom_vision_model = CustomVisionModel(PREDICTION_KEY, ENDPOINT, PROJECT_ID, ITERATION_NAME)
phi_model = PhiModel(MODEL_CATALOG_ENDPOINT, MODEL_CATALOG_API_KEY)
llama_model = LlamaModel(LLAMA_ENDPOINT, LLAMA_TOKEN, LLAMA_MODEL_NAME)

# Initializing handlers for Custom Vision, Feedback, and Llama models
# Custom Vision, Phi 및 Llama 모델 핸들러 초기화
custom_vision_handler = CustomVisionHandler(PREDICTION_KEY, ENDPOINT, PROJECT_ID, ITERATION_NAME)
feedback_handler = FeedbackHandler(MODEL_CATALOG_ENDPOINT, MODEL_CATALOG_API_KEY, LLAMA_ENDPOINT, LLAMA_TOKEN, LLAMA_MODEL_NAME, translation_handler)

# Function to display feedback block on the screen
# 피드백 블럭으로 화면 출력 함수
def display_feedback(feedback, language):
    expander_title = translation_handler.translate_text("🔍 Combined Model Feedback", language)
    info_title = translation_handler.translate_text("**Combined Model Feedback**", language)
    
    with st.expander(expander_title, expanded=True):
        st.info(info_title, icon="💡")
        st.write(feedback)

# Function to create .docx file for feedback
# 피드백을 위한 .docx 파일 생성 함수
def create_docx_feedback(feedback_text, filename="Feedback.docx"):
    doc = Document()
    doc.add_heading("Model Feedback", level=1)
    doc.add_paragraph(feedback_text)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Running the app
# 앱 실행 함수
def run_app():
    # Adding language selection in the sidebar
    # 사이드바에 언어 선택 추가
    language = st.selectbox("Select language", ["EN", "KO", "ZH", "JA", "ES"])
    st.title(translation_handler.translate_text("AI-based Skin Disease Diagnosis Website", language))

    # Translating user input options
    # 사용자 입력 옵션을 번역
    gender_options = [translation_handler.translate_text("Male", language), translation_handler.translate_text("Female", language)]
    age_group_options = [translation_handler.translate_text("Under 20", language), translation_handler.translate_text("20-29", language), translation_handler.translate_text("30-39", language),
                         translation_handler.translate_text("40-49", language), translation_handler.translate_text("50-59", language), translation_handler.translate_text("60-69", language), 
                         translation_handler.translate_text("70 and above", language)]
    city_options = [translation_handler.translate_text("New York", language), translation_handler.translate_text("London", language), translation_handler.translate_text("Tokyo", language), 
                    translation_handler.translate_text("Paris", language), translation_handler.translate_text("Beijing", language), translation_handler.translate_text("Sydney", language), 
                    translation_handler.translate_text("Seoul", language)]

    # Receiving user input in translated form
    # 번역된 형태로 사용자 입력 받기
    gender = st.selectbox(translation_handler.translate_text("Select your gender", language), gender_options)
    age_group = st.selectbox(translation_handler.translate_text("Select your age group", language), age_group_options)
    city = st.selectbox(translation_handler.translate_text("Select your City", language), city_options)

    # Converting user input back to English
    # 사용자 입력값을 영어로 다시 변환
    gender_en = translation_handler.translate_text(gender, "EN")
    age_group_en = translation_handler.translate_text(age_group, "EN")
    city_en = translation_handler.translate_text(city, "EN")

    # File upload for image
    # 이미지 파일 업로드
    uploaded_file = st.file_uploader(translation_handler.translate_text("Upload an image", language), type=["jpg", "jpeg", "png", "bmp"])

    if uploaded_file is not None:
        image = Image.open(uploaded_file)  # Open the uploaded image
        st.image(image, caption=translation_handler.translate_text("Uploaded Image", language), use_column_width=True)  # Display the image
        image_data = custom_vision_handler.convert_image_to_bytes(image)  # Convert image to byte format

        # If "Predict" button is clicked, save prediction result in session state
        # "Predict" 버튼 클릭 시 예측 결과를 세션 상태에 저장
        if st.button(translation_handler.translate_text("Predict", language)):
            predictions = custom_vision_handler.predict_image(image_data)  # Run prediction
            if predictions:
                top_prediction = predictions[0]  # Get the top prediction result
                prediction_text = (
                    f"The probability that this image represents a skin condition called "
                    f"{top_prediction.tag_name} is {top_prediction.probability * 100:.2f}%."
                )
                st.session_state.prediction_text = translation_handler.translate_text(prediction_text, language)
                st.session_state.prediction_tag = top_prediction.tag_name
            else:
                st.session_state.prediction_text = translation_handler.translate_text("No prediction found.", language)
                st.session_state.prediction_tag = None

        # Display prediction result only if available
        # 예측 결과가 있을 때만 출력
        if "prediction_text" in st.session_state:
            st.write(st.session_state.prediction_text)

        # If prediction result exists and feedback button is clicked, generate feedback
        # 예측 결과가 있고 피드백 버튼이 눌리면 피드백 생성
        if st.session_state.get("prediction_tag") and st.button(translation_handler.translate_text("Model Feedback", language)):
            prompt_1 = f"Recommend treatment for a {age_group_en} {gender_en} in {city_en} with a skin condition called {st.session_state.prediction_tag}."
            prompt_2 = f"For a {age_group_en} {gender_en} with a skin condition called {st.session_state.prediction_tag}, please recommend hospitals in {city_en}."

            feedback = feedback_handler.generate_feedback(prompt_1, prompt_2, language)  # Generate feedback
            feedback_handler.display_feedback(feedback, language)  # Display feedback on screen

            # Create and provide download button for .docx file
            # .docx 파일을 생성하고 다운로드 버튼을 제공
            file_stream = create_docx_feedback(feedback, filename="Feedback.docx")

            st.download_button(
                label=translation_handler.translate_text("Download Feedback as DOCX", language),
                data=file_stream,
                file_name="Feedback.docx",  # Set DOCX file name
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # DOCX MIME type
            )

# Run the app
# 앱 실행
run_app()
